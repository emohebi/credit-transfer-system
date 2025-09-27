import json
import re
from docx import Document
from typing import Dict, List, Optional

def extract_vet_course_info(docx_path: str, debug: bool = False) -> Dict:
    """
    Extract VET course information from a DOCX file and format as JSON.
    
    Args:
        docx_path: Path to the DOCX file
        debug: If True, print debug information
        
    Returns:
        Dictionary containing course information
    """
    # Read the document
    doc = Document(docx_path)
    
    # Extract all text with paragraph tracking
    full_text = []
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
            paragraphs.append(text)
    
    # Also extract text from tables
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_data.append(row_text)
    
    # Initialize result dictionary
    result = {
        "code": None,
        "name": None,
        "description": None,
        "learning_outcomes": [],
        "assessment_requirements": None,
        "nominal_hours": None,
        "prerequisites": []
    }
    
    # Extract code and name
    for i, para in enumerate(paragraphs):
        # Look for pattern like "BSBOPS501 Manage business resources"
        match = re.match(r'^([A-Z]+\d+)\s+(.+)$', para)
        if match:
            result["code"] = match.group(1)
            result["name"] = match.group(2)
            break
    
    # Extract sections - improved section detection
    sections = {}
    current_section = None
    current_content = []
    
    for i, para in enumerate(paragraphs):
        # Check if this is a section header (multiple detection methods)
        is_header = False
        cleaned_para = para.replace('#', '').strip()
        
        # Method 1: Starts with #
        if para.startswith('#'):
            is_header = True
            current_section = cleaned_para
        # Method 2: Exact match of known section names
        elif cleaned_para in ['Application', 'Elements and Performance Criteria', 
                            'Foundation Skills', 'Performance Evidence', 
                            'Knowledge Evidence', 'Assessment Conditions',
                            'Unit Mapping Information', 'Unit Sector',
                            'Modification History', 'Links']:
            is_header = True
            current_section = cleaned_para
        # Method 3: Check if line is a header-like format (all caps or title case main sections)
        elif (len(para) < 50 and 
              (para.isupper() or 
               para in ['Performance Evidence', 'Knowledge Evidence', 'Assessment Conditions'] or
               'Evidence' in para or 'Assessment' in para or 'Application' in para)):
            # Additional check to avoid false positives
            if i + 1 < len(paragraphs) and not paragraphs[i + 1].startswith(para):
                is_header = True
                current_section = cleaned_para
        
        if is_header:
            # Save previous section
            if current_section and current_content:
                sections[current_section] = '\n'.join(current_content)
            current_content = []
        else:
            if current_section:
                current_content.append(para)
    
    # Add the last section
    if current_section:
        sections[current_section] = '\n'.join(current_content)
    
    # Debug: Print available sections
    if debug:
        print("Available sections found:")
        for section_name in sections.keys():
            print(f"  - {section_name}")
            print(f"    Content preview: {sections[section_name][:100]}...")
    
    # Build description combining Application, Elements and Performance Criteria, and Foundation Skills
    description_parts = []
    
    # Add Application section
    if 'Application' in sections:
        description_parts.append("Application:\n" + sections['Application'])
    
    # If Application not found in sections, search for it in the full text
    if not description_parts:
        full_doc_text = '\n'.join(paragraphs)
        app_match = re.search(r'Application\s*\n+(.*?)(?=Unit Sector|Elements and Performance Criteria|Foundation Skills|#|$)', 
                             full_doc_text, re.DOTALL)
        if app_match:
            app_content = app_match.group(1).strip()
            if app_content:
                description_parts.append("Application:\n" + app_content)
    
    # Add Elements and Performance Criteria - extract ALL from table data
    if 'Elements and Performance Criteria' in sections or table_data:
        perf_criteria = []
        in_perf_table = False
        for row in table_data:
            if len(row) >= 2:
                # Check if this is the performance criteria table
                if 'ELEMENT' in row[0] or 'PERFORMANCE CRITERIA' in row[1]:
                    in_perf_table = True
                    continue
                if in_perf_table and row[0] and row[1]:
                    # Skip header rows
                    if not row[0].startswith('*'):
                        perf_criteria.append(f"{row[0]}: {row[1]}")
        
        if perf_criteria:
            description_parts.append("\n\nElements and Performance Criteria:\n" + '\n'.join(perf_criteria))
    
    # Add ALL Foundation Skills
    if 'Foundation Skills' in sections:
        foundation_skills = []
        for row in table_data:
            if len(row) >= 2:
                skill_names = ['Reading', 'Writing', 'Oral communication', 'Numeracy', 
                              'Enterprise and initiative', 'Teamwork', 'Planning and organising']
                if row[0] in skill_names:
                    foundation_skills.append(f"{row[0]}: {row[1]}")
        
        if foundation_skills:
            description_parts.append("\n\nFoundation Skills:\n" + '\n'.join(foundation_skills))
    
    result["description"] = '\n'.join(description_parts)
    
    # Extract ALL learning outcomes from Knowledge Evidence
    learning_outcomes = []
    
    # Method 1: Check if Knowledge Evidence is in sections
    if 'Knowledge Evidence' in sections:
        knowledge_text = sections['Knowledge Evidence']
        
        # Parse the Knowledge Evidence text which typically contains bullet points
        lines = knowledge_text.split('\n')
        for line in lines:
            line = line.strip()
            if line and not line.startswith('The candidate must'):
                # Remove bullet point markers
                clean_line = re.sub(r'^[-•·]\s*', '', line).strip()
                if clean_line:
                    learning_outcomes.append(clean_line)
    
    # Method 2: If not found in sections, search in the full text
    if not learning_outcomes:
        # Find Knowledge Evidence section in full text
        full_doc_text = '\n'.join(paragraphs)
        knowledge_match = re.search(r'Knowledge Evidence(.*?)(?=Assessment Conditions|Assessment Requirements|Links|$)', 
                                   full_doc_text, re.DOTALL)
        if knowledge_match:
            knowledge_content = knowledge_match.group(1)
            
            # Extract items that look like knowledge points
            lines = knowledge_content.split('\n')
            for line in lines:
                line = line.strip()
                # Skip the preamble and empty lines
                if (line and 
                    not line.startswith('The candidate must') and
                    not line == 'Knowledge Evidence'):
                    # Clean up bullet points and add to outcomes
                    clean_line = re.sub(r'^[-•·]\s*', '', line).strip()
                    if clean_line and len(clean_line) > 5:
                        learning_outcomes.append(clean_line)
    
    # Method 3: Extract from table data if Knowledge Evidence is in a table
    if not learning_outcomes:
        for i, row in enumerate(table_data):
            # Check if we're in a Knowledge Evidence section
            if any('Knowledge Evidence' in cell for cell in row):
                # Start extracting from next rows
                for j in range(i+1, len(table_data)):
                    if table_data[j]:
                        for cell in table_data[j]:
                            if cell and not cell.startswith('Assessment'):
                                learning_outcomes.append(cell)
                break
    
    # Remove duplicates while preserving order
    seen = set()
    unique_outcomes = []
    for outcome in learning_outcomes:
        if outcome not in seen:
            seen.add(outcome)
            unique_outcomes.append(outcome)
    
    result["learning_outcomes"] = unique_outcomes
    
    if debug and not unique_outcomes:
        print("Warning: No learning outcomes extracted from Knowledge Evidence")
        print("Sections available:", list(sections.keys()))
    
    # Extract COMPLETE assessment requirements from Assessment Conditions and Performance Evidence
    assessment_parts = []
    
    if 'Performance Evidence' in sections:
        perf_evidence = sections['Performance Evidence']
        # Add the complete performance evidence section
        assessment_parts.append("Performance Evidence: " + perf_evidence)
    
    if 'Assessment Conditions' in sections:
        assessment_cond = sections['Assessment Conditions']
        # Add the complete assessment conditions section
        assessment_parts.append("Assessment Conditions: " + assessment_cond)
    
    result["assessment_requirements"] = '\n\n'.join(assessment_parts) if assessment_parts else None
    
    # Extract prerequisites from Unit Mapping Information
    if 'Unit Mapping Information' in sections:
        mapping_info = sections['Unit Mapping Information']
        # Look for prerequisite patterns
        if 'prerequisite' in mapping_info.lower():
            # Extract ALL unit codes that might be prerequisites
            prereq_codes = re.findall(r'\b[A-Z]+\d+\b', mapping_info)
            result["prerequisites"] = prereq_codes
        else:
            # If no explicit prerequisites mentioned, leave empty
            result["prerequisites"] = []
    
    # Nominal hours would need to be extracted from additional documentation
    # Setting to None as it's not in this document
    result["nominal_hours"] = None
    
    return result

def save_to_json(data: Dict, output_path: str = "vet_course.json"):
    """
    Save the extracted data to a JSON file.
    
    Args:
        data: Dictionary containing course information
        output_path: Path for the output JSON file
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"JSON data saved to {output_path}")

def main():
    """
    Main function to run the extraction.
    """
    qual = {
    "code": "BSB50120",
    "name": "Diploma of Business",
    "level": "Diploma",
    "units": []
    }
    # Input file path
    docx_file = "/home/ehsan/Downloads/BSBOPS501_Complete_R1.docx"  # Update with your actual file path
    
    try:
        # Extract course information
        course_info = extract_vet_course_info(docx_file)
        qual["units"].append(course_info)
        # Print the result
        # print("Extracted Course Information:")
        # print(json.dumps(course_info, indent=2, ensure_ascii=False))
        
        # Save to JSON file
        save_to_json(qual, output_path="./data/diploma_of_business.json")
        
    except FileNotFoundError:
        print(f"Error: File '{docx_file}' not found.")
        print("Please ensure the DOCX file is in the same directory as this script.")
    except Exception as e:
        print(f"Error processing file: {e}")
        print("Please ensure you have python-docx installed: pip install python-docx")

if __name__ == "__main__":
    main()